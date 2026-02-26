from fastapi import FastAPI, UploadFile, File
import shutil, os, time
import fitz  # PyMuPDF
import numpy as np
from PIL import Image

from api.ocr_engine import extract_single_page
from api.parsers import detect_doc_type
from api.ner_parser import extract_with_ner

app = FastAPI(title="OCR Extraction API", version="6.0.0", docs_url="/docs")

UPLOAD_DIR = "temp"
os.makedirs(UPLOAD_DIR, exist_ok=True)
SCANNED_THRESHOLD = 50

SUPPORTED = (".pdf", ".png", ".jpg", ".jpeg", ".docx", ".doc")


# ── PyMuPDF table → items ─────────────────────────────────────────────────────
def pymupdf_table_to_items(page) -> list:
    items = []
    try:
        finder = page.find_tables()
        for tbl in finder.tables:
            df = tbl.to_pandas()
            if df.empty or len(df) < 2:
                continue
            for _, row in df.iterrows():
                item = {str(k).strip(): str(v).strip()
                        for k, v in row.items()
                        if str(v) not in ("nan", "None", "")}
                vals_lower = str(list(item.values())).lower()
                if any(h in vals_lower for h in ["description","item","hsn","qty","amount","rate"]):
                    continue
                if item:
                    items.append(item)
    except Exception:
        pass
    return items


# ── DOCX text extractor ───────────────────────────────────────────────────────
def extract_docx_text(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        lines = []
        for para in doc.paragraphs:
            if para.text.strip():
                lines.append(para.text.strip())
        # Extract tables too
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    lines.append(row_text)
        return "\n".join(lines)
    except Exception as e:
        return f"Error reading docx: {e}"


# ── PDF text extractor ────────────────────────────────────────────────────────
def extract_pdf_pages(file_path: str):
    doc = fitz.open(file_path)
    pages_output = []
    confidence = 1.0

    for page_num, page in enumerate(doc, start=1):
        native_text = page.get_text("text").strip()
        is_scanned  = len(native_text) < SCANNED_THRESHOLD

        if not is_scanned:
            page_text  = native_text
            confidence = 1.0
        else:
            mat = fitz.Matrix(150/72, 150/72)
            pix = page.get_pixmap(matrix=mat)
            import cv2
            arr     = np.frombuffer(pix.tobytes("png"), dtype=np.uint8)
            img     = cv2.imdecode(arr, cv2.IMREAD_COLOR)
            pil_img = Image.fromarray(cv2.cvtColor(img, cv2.COLOR_BGR2RGB))
            res        = extract_single_page(pil_img)
            page_text  = res["formatted_text"]
            confidence = res["confidence_score"]

        doc_type = detect_doc_type(page_text)
        fields   = extract_with_ner(page_text, doc_type)

        if doc_type in ("invoice", "purchase_order") and not is_scanned:
            pymupdf_items = pymupdf_table_to_items(page)
            if pymupdf_items:
                fields["items"] = pymupdf_items

        pages_output.append({
            "page":     page_num,
            "doc_type": doc_type,
            "fields":   fields,
            "text":     page_text,
        })

    doc.close()
    return pages_output, round(confidence, 3)


# ── Universal /extract endpoint ───────────────────────────────────────────────
@app.post("/extract")
async def extract_any(file: UploadFile = File(...)):
    fname = file.filename.lower()
    ext   = os.path.splitext(fname)[1]

    if ext not in SUPPORTED:
        return {"error": f"Unsupported format: {ext}. Supported: {', '.join(SUPPORTED)}"}

    file_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(file_path, "wb") as buffer:
        shutil.copyfileobj(file.file, buffer)

    start = time.time()

    # ── PDF ──
    if ext == ".pdf":
        pages_output, confidence = extract_pdf_pages(file_path)
        elapsed = round(time.time() - start, 2)
        return {
            "status":      "success",
            "file_type":   "pdf",
            "total_pages": len(pages_output),
            "confidence":  confidence,
            "pages":       pages_output,
            "meta": {
                "file_name":           file.filename,
                "processing_time_sec": elapsed,
                "extraction_method":   "spacy_ner + pymupdf",
            },
        }

    # ── IMAGE ──
    elif ext in (".png", ".jpg", ".jpeg"):
        image    = Image.open(file_path)
        result   = extract_single_page(image)
        text     = result["formatted_text"]
        doc_type = detect_doc_type(text)
        fields   = extract_with_ner(text, doc_type)
        elapsed  = round(time.time() - start, 2)
        return {
            "status":     "success",
            "file_type":  "image",
            "doc_type":   doc_type,
            "confidence": round(result["confidence_score"], 3),
            "fields":     fields,
            "raw_text":   text,
            "meta": {
                "file_name":           file.filename,
                "processing_time_sec": elapsed,
                "extraction_method":   "spacy_ner + paddleocr",
            },
        }

    # ── WORD DOCUMENT ──
    elif ext in (".docx", ".doc"):
        text     = extract_docx_text(file_path)
        doc_type = detect_doc_type(text)
        fields   = extract_with_ner(text, doc_type)
        elapsed  = round(time.time() - start, 2)
        return {
            "status":     "success",
            "file_type":  "word",
            "doc_type":   doc_type,
            "confidence": 1.0,
            "fields":     fields,
            "raw_text":   text,
            "meta": {
                "file_name":           file.filename,
                "processing_time_sec": elapsed,
                "extraction_method":   "spacy_ner + python-docx",
            },
        }


# ── Keep old endpoints for backward compatibility ─────────────────────────────
@app.post("/extract/pdf")
async def extract_pdf_api(file: UploadFile = File(...)):
    return await extract_any(file)

@app.post("/extract/image")
async def extract_image_api(file: UploadFile = File(...)):
    return await extract_any(file)


# ── Health check ──────────────────────────────────────────────────────────────
@app.get("/health")
def health():
    return {
        "status":          "healthy",
        "version":         "6.0.0",
        "model":           "spaCy NER (ocr_ner_model)",
        "mode":            "offline",
        "supported_formats": list(SUPPORTED),
    }